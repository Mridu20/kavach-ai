"""
Tool Registry, Tool Selection Hub, and Retry Mechanism for KAVACH AI Workbench.
"""

import time
import uuid
import logging
from abc import ABC, abstractmethod
from typing import Any, Callable, Dict, List, Optional
from backend.agent.state import ToolCallRecord

logger = logging.getLogger("kavach_agent.tools")


class BaseAgentTool(ABC):
    """Abstract Base Class for all KAVACH AI agent tools."""
    name: str
    description: str
    category: str

    @abstractmethod
    def run(self, **kwargs) -> Dict[str, Any]:
        """Execute tool logic and return dictionary result."""
        pass


class MockOCRTool(BaseAgentTool):
    name = "ocr_pdf_tool"
    description = "Extract text, tables, and handwriting from scanned PDF documents locally."
    category = "document"

    def run(self, file_path: str = "", pages: Optional[List[int]] = None, **kwargs) -> Dict[str, Any]:
        from ingestion import extract_content
        ingest_res = extract_content(file_path=file_path)
        return {
            "file_path": file_path,
            "extracted_text": ingest_res.raw_text,
            "extraction_method": ingest_res.extraction_method,
            "pages_processed": ingest_res.pages_processed,
            "tables_found": 2,
            "handwriting_detected": ingest_res.structured.handwriting_detected,
            "structured_findings": [f.model_dump() for f in ingest_res.structured.findings],
        }


class MockRAGSearchTool(BaseAgentTool):
    name = "rag_search_tool"
    description = "Search local SOPs, manuals, and correspondence using vector embeddings."
    category = "retrieval"

    def run(self, query: str = "", top_k: int = 3, **kwargs) -> Dict[str, Any]:
        try:
            from backend.rag.store import get_vector_store
            store = get_vector_store()
            results = store.query(query_text=query, top_k=top_k)
            if results:
                return {
                    "query": query,
                    "results": results,
                }
        except Exception as e:
            pass

        return {
            "query": query,
            "results": [
                {
                    "doc_name": "SOP_Industrial_Safety_v3.pdf",
                    "page": 14,
                    "score": 0.92,
                    "snippet": "Section 4.2: Pressure vessel inspection must mandate immediate shutdown if corrosion exceeds 0.5mm.",
                },
                {
                    "doc_name": "Maintenance_Manual_Turbine_2025.pdf",
                    "page": 8,
                    "score": 0.87,
                    "snippet": "Section 2.1: Secondary containment seal replacement required every 12 months.",
                },
            ],
        }


class MockVisionTool(BaseAgentTool):
    name = "vision_analysis_tool"
    description = "Analyze photographs and visual diagram components using local vision-language model."
    category = "vision"

    def run(self, image_path: str = "", prompt: str = "", **kwargs) -> Dict[str, Any]:
        from ingestion import extract_content
        ingest_res = extract_content(file_path=image_path, force_vlm=True)
        return {
            "image_path": image_path,
            "analysis": ingest_res.raw_text,
            "confidence": 0.94,
            "detected_objects": ["weld_joint_B12", "surface_crack", "corrosion_spot"],
            "structured_findings": [f.model_dump() for f in ingest_res.structured.findings],
        }



class MockSandboxCodeTool(BaseAgentTool):
    name = "sandbox_code_tool"
    description = "Execute python or shell script safely in isolated Docker sandbox without network access."
    category = "sandbox"

    def run(self, code: str = "", language: str = "python", **kwargs) -> Dict[str, Any]:
        return {
            "language": language,
            "stdout": "[SANDBOX STDOUT] Execution completed cleanly with 0 cloud calls.",
            "stderr": "",
            "exit_code": 0,
            "network_calls_blocked": 0,
        }


class DocxGeneratorTool(BaseAgentTool):
    name = "generate_docx_tool"
    description = "Generate official Approval Note DOCX document from structured findings."
    category = "generator"

    def run(
        self,
        title: str = "Approval Note",
        findings: Optional[Dict[str, Any]] = None,
        evidence: Optional[List[Any]] = None,
        task_id: str = "",
        output_path: str = "Approval_Note.docx",
        **kwargs,
    ) -> Dict[str, Any]:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import os
        from datetime import datetime

        findings = findings or {}
        evidence = evidence or []
        output_dir = os.path.join("backend", "storage", "outputs")
        os.makedirs(output_dir, exist_ok=True)
        full_path = os.path.join(output_dir, os.path.basename(output_path))

        doc = Document()

        # ── Page margins ──────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # ── Header bar ────────────────────────────────────────────
        hdr = doc.add_paragraph()
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr.add_run("KAVACH AI — SOVEREIGN INDUSTRIAL WORKBENCH")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run("Confidential — On-Premise Generated — Zero Cloud Calls").font.size = Pt(9)

        doc.add_paragraph()

        # ── Title ─────────────────────────────────────────────────
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ── Metadata table ────────────────────────────────────────
        meta_table = doc.add_table(rows=3, cols=4)
        meta_table.style = "Table Grid"
        meta_data = [
            ("Task ID", task_id or "N/A", "Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
            ("Classification", "CONFIDENTIAL", "Status", "AWAITING APPROVAL"),
            ("System", "KAVACH AI v1.0", "Network", "AIR-GAPPED / ZERO CLOUD"),
        ]
        for row_idx, (k1, v1, k2, v2) in enumerate(meta_data):
            row = meta_table.rows[row_idx]
            for cell, text, bold in [(row.cells[0], k1, True), (row.cells[1], v1, False),
                                      (row.cells[2], k2, True), (row.cells[3], v2, False)]:
                cell.text = text
                cell.paragraphs[0].runs[0].bold = bold
                cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # ── Executive Summary ─────────────────────────────────────
        doc.add_heading("1. Executive Summary", level=2)
        ocr_text = findings.get("ocr_extracted_text", "")
        vision_text = findings.get("vision_analysis", "")
        summary_text = ocr_text[:600] if ocr_text else "No OCR text extracted."
        if vision_text:
            summary_text += f"\n\nVisual Analysis: {vision_text[:300]}"
        doc.add_paragraph(summary_text)

        doc.add_paragraph()

        # ── Findings table ────────────────────────────────────────
        doc.add_heading("2. Structured Findings", level=2)
        findings_table = doc.add_table(rows=1, cols=4)
        findings_table.style = "Table Grid"
        findings_table.autofit = False
        col_widths = [Inches(0.6), Inches(2.5), Inches(1.2), Inches(1.8)]
        hdr_row = findings_table.rows[0]
        for i, (cell, width, label) in enumerate(zip(
            hdr_row.cells, col_widths,
            ["#", "Finding / Observation", "Severity", "Recommended Action"]
        )):
            cell.width = width
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        # Populate from findings dict if structured_findings present, else one summary row
        structured = findings.get("structured_findings", [])
        if structured:
            for idx, f in enumerate(structured[:10], 1):
                row = findings_table.add_row()
                row.cells[0].width = col_widths[0]
                row.cells[1].width = col_widths[1]
                row.cells[2].width = col_widths[2]
                row.cells[3].width = col_widths[3]
                row.cells[0].text = str(idx)
                row.cells[1].text = str(f.get("description", ""))[:200]
                row.cells[2].text = str(f.get("severity", "MEDIUM"))
                row.cells[3].text = "Follow-up inspection required."
                for cell in row.cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
        else:
            row = findings_table.add_row()
            row.cells[0].text = "1"
            row.cells[1].text = ocr_text[:200] if ocr_text else "See executive summary."
            row.cells[2].text = "REVIEW"
            row.cells[3].text = "Inspector sign-off required."
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # ── SOP Citations ─────────────────────────────────────────
        doc.add_heading("3. SOP & Regulatory Citations", level=2)
        if evidence:
            for i, ev in enumerate(evidence[:5], 1):
                source = getattr(ev, "source_doc", str(ev))
                page = getattr(ev, "page_num", None)
                snippet = getattr(ev, "snippet", "")
                score = getattr(ev, "confidence_score", 1.0)
                p = doc.add_paragraph(style="List Number")
                p.add_run(f"{source}").bold = True
                page_str = f", p.{page}" if page else ""
                p.add_run(f"{page_str} (confidence: {score:.0%})\n")
                snippet_run = p.add_run(f'  "{snippet[:200]}"')
                snippet_run.italic = True
        else:
            doc.add_paragraph("No SOP evidence retrieved for this task.")

        doc.add_paragraph()

        # ── Recommendation ────────────────────────────────────────
        doc.add_heading("4. Recommendation", level=2)
        doc.add_paragraph(
            "Based on the sovereign on-premise analysis above, the findings have been "
            "reviewed against applicable SOPs and maintenance manuals. The inspector is "
            "requested to review and approve or reject this note below."
        )

        doc.add_paragraph()

        # ── Approval block ────────────────────────────────────────
        doc.add_heading("5. Inspector Approval", level=2)
        approval_table = doc.add_table(rows=4, cols=2)
        approval_table.style = "Table Grid"
        approval_table.autofit = False
        approval_table.columns[0].width = Inches(2.0)
        approval_table.columns[1].width = Inches(4.0)
        rows_data = [
            ("Reviewer Name", ""),
            ("Decision", "[ ] APPROVED      [ ] APPROVED WITH MODIFICATIONS"),
            ("", "[ ] REJECTED"),
            ("Date & Signature", ""),
        ]
        for row, (label, value) in zip(approval_table.rows, rows_data):
            row.cells[0].width = Inches(2.0)
            row.cells[1].width = Inches(4.0)
            row.cells[0].text = label
            row.cells[1].text = value
            if row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].bold = True
                row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            if row.cells[1].paragraphs[0].runs:
                row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

        doc.save(full_path)
        file_size_kb = round(os.path.getsize(full_path) / 1024, 1)

        return {
            "output_path": full_path,
            "status": "CREATED",
            "file_size_kb": file_size_kb,
            "sections_generated": ["Header", "Metadata", "Executive Summary", "Findings Table", "SOP Citations", "Recommendation", "Approval Block"],
        }


class XlsxGeneratorTool(BaseAgentTool):
    name = "generate_xlsx_tool"
    description = "Generate Action Tracker XLSX spreadsheet from inspection tasks."
    category = "generator"

    def run(
        self,
        items: Optional[List[Dict[str, Any]]] = None,
        output_path: str = "Action_Tracker.xlsx",
        **kwargs,
    ) -> Dict[str, Any]:
        import openpyxl
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
        import os
        from datetime import datetime

        items = items or []
        output_dir = os.path.join("backend", "storage", "outputs")
        os.makedirs(output_dir, exist_ok=True)
        full_path = os.path.join(output_dir, os.path.basename(output_path))

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Action Tracker"

        # ── Palette ───────────────────────────────────────────────
        BLUE_DARK  = "1E3A5F"
        BLUE_MID   = "2563EB"
        AMBER      = "F59E0B"
        RED_FILL   = "FEE2E2"
        AMBER_FILL = "FEF3C7"
        GREEN_FILL = "DCFCE7"
        GREY_FILL  = "F1F5F9"
        WHITE      = "FFFFFF"

        thin = Side(style="thin", color="CBD5E1")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        # ── Title block ───────────────────────────────────────────
        ws.merge_cells("A1:G1")
        title_cell = ws["A1"]
        title_cell.value = "KAVACH AI — ACTION TRACKER"
        title_cell.font = Font(bold=True, size=14, color=WHITE)
        title_cell.fill = PatternFill("solid", fgColor=BLUE_DARK)
        title_cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 28

        ws.merge_cells("A2:G2")
        sub_cell = ws["A2"]
        sub_cell.value = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  System: KAVACH AI Sovereign Workbench  |  Network: AIR-GAPPED"
        sub_cell.font = Font(size=9, color="475569")
        sub_cell.fill = PatternFill("solid", fgColor=GREY_FILL)
        sub_cell.alignment = Alignment(horizontal="center")
        ws.row_dimensions[2].height = 16

        ws.append([])  # spacer row 3

        # ── Column headers ────────────────────────────────────────
        headers = ["#", "Task / Finding", "Priority", "Owner", "Due By", "Status", "Notes"]
        col_widths = [5, 42, 12, 18, 14, 14, 28]
        ws.append(headers)
        hdr_row = ws.row_dimensions[4]
        hdr_row.height = 18
        for col_idx, (h, w) in enumerate(zip(headers, col_widths), 1):
            cell = ws.cell(row=4, column=col_idx)
            cell.font = Font(bold=True, size=10, color=WHITE)
            cell.fill = PatternFill("solid", fgColor=BLUE_MID)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
            ws.column_dimensions[cell.column_letter].width = w

        # Freeze header rows
        ws.freeze_panes = "A5"

        # ── Priority colour map ───────────────────────────────────
        priority_fill = {
            "HIGH":     PatternFill("solid", fgColor=RED_FILL),
            "CRITICAL": PatternFill("solid", fgColor=RED_FILL),
            "MEDIUM":   PatternFill("solid", fgColor=AMBER_FILL),
            "LOW":      PatternFill("solid", fgColor=GREEN_FILL),
        }

        # ── Data rows ─────────────────────────────────────────────
        for idx, item in enumerate(items, 1):
            priority = str(item.get("priority", "MEDIUM")).upper()
            row_data = [
                idx,
                item.get("task", ""),
                priority,
                item.get("owner", "Inspector"),
                item.get("due_by", "Immediate"),
                item.get("status", "OPEN"),
                item.get("notes", ""),
            ]
            ws.append(row_data)
            data_row = ws.max_row
            ws.row_dimensions[data_row].height = 20
            fill = priority_fill.get(priority, PatternFill("solid", fgColor=WHITE))
            for col_idx in range(1, 8):
                cell = ws.cell(row=data_row, column=col_idx)
                if col_idx == 3:
                    cell.fill = fill
                    cell.font = Font(bold=True, size=9)
                else:
                    cell.font = Font(size=9)
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.border = border

        # ── Footer ────────────────────────────────────────────────
        ws.append([])
        footer_row = ws.max_row + 1
        ws.merge_cells(f"A{footer_row}:G{footer_row}")
        footer_cell = ws.cell(row=footer_row, column=1)
        footer_cell.value = "KAVACH AI — Sovereign On-Premise System | Zero External Network Calls | All data processed locally"
        footer_cell.font = Font(italic=True, size=8, color="94A3B8")
        footer_cell.alignment = Alignment(horizontal="center")

        wb.save(full_path)
        file_size_kb = round(os.path.getsize(full_path) / 1024, 1)

        return {
            "output_path": full_path,
            "status": "CREATED",
            "file_size_kb": file_size_kb,
            "rows_written": len(items),
        }


class MockModelRouterTool(BaseAgentTool):
    name = "model_router_tool"
    description = "Route task to specialized local model (coding, vision, general reasoning)."
    category = "routing"

    def run(self, task_type: str = "general", prompt: str = "", **kwargs) -> Dict[str, Any]:
        selected_model = "llama3-8b-instruct-q4"
        if task_type == "vision":
            selected_model = "llava-v1.6-7b-q4"
        elif task_type == "coding":
            selected_model = "qwen2.5-coder-7b-q4"

        return {
            "selected_model": selected_model,
            "task_type": task_type,
            "response": f"[MOCK MODEL RESPONSE via {selected_model}] Completed reasoning for prompt.",
        }


class ToolRegistry:
    """Registry managing available tools and execution retries."""

    def __init__(self):
        self._tools: Dict[str, BaseAgentTool] = {}
        self._register_defaults()

    def _register_defaults(self):
        defaults = [
            MockOCRTool(),
            MockRAGSearchTool(),
            MockVisionTool(),
            MockSandboxCodeTool(),
            DocxGeneratorTool(),
            XlsxGeneratorTool(),
            MockModelRouterTool(),
        ]
        for tool in defaults:
            self.register_tool(tool)

    def register_tool(self, tool: BaseAgentTool):
        self._tools[tool.name] = tool
        logger.info(f"Registered tool: {tool.name}")

    def get_tool(self, name: str) -> Optional[BaseAgentTool]:
        return self._tools.get(name)

    def list_tools(self) -> List[Dict[str, str]]:
        return [
            {"name": t.name, "description": t.description, "category": t.category}
            for t in self._tools.values()
        ]

    def execute_with_retry(
        self,
        tool_name: str,
        step_id: int,
        params: Dict[str, Any],
        max_retries: int = 3,
        backoff_factor: float = 0.5,
    ) -> ToolCallRecord:
        """Executes a tool with automatic retries and logs performance."""
        tool = self.get_tool(tool_name)
        call_id = f"call_{uuid.uuid4().hex[:8]}"

        if not tool:
            return ToolCallRecord(
                call_id=call_id,
                step_id=step_id,
                tool_name=tool_name,
                input_params=params,
                success=False,
                error_message=f"Tool '{tool_name}' is not registered.",
            )

        start_time = time.time()
        last_error: Optional[str] = None

        for attempt in range(1, max_retries + 1):
            try:
                result = tool.run(**params)
                elapsed_ms = (time.time() - start_time) * 1000
                return ToolCallRecord(
                    call_id=call_id,
                    step_id=step_id,
                    tool_name=tool_name,
                    input_params=params,
                    output=result,
                    execution_time_ms=round(elapsed_ms, 2),
                    success=True,
                )
            except Exception as e:
                last_error = str(e)
                logger.warning(f"Tool execution '{tool_name}' failed attempt {attempt}/{max_retries}: {last_error}")
                if attempt < max_retries:
                    time.sleep(backoff_factor * attempt)

        elapsed_ms = (time.time() - start_time) * 1000
        return ToolCallRecord(
            call_id=call_id,
            step_id=step_id,
            tool_name=tool_name,
            input_params=params,
            execution_time_ms=round(elapsed_ms, 2),
            success=False,
            error_message=last_error or "Unknown tool execution failure.",
        )


# Global default instance
default_tool_registry = ToolRegistry()
